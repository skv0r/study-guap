import os
import glob
from docx import Document
from docx.shared import Pt, Inches

def parse_txt(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    data = {"goal": "", "task": "", "conclusion": "", "title1": "", "title2": ""}
    sec = None
    for l in lines:
        s = l.strip()
        if "ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №" in s:
            data["title1"] = s
        elif data["title1"] and not data["title2"] and s and not any(x in s for x in ["ВЫПОЛНИЛ", "ПРЕПОДАВАТЕЛЬ", "Санкт-Петербург"]):
            if "РАБОТА С" in s or "ФОРМЫ" in s or "ВЕРСТКА" in s or "ГИПЕР" in s:
                data["title2"] = s
        
        if "1. Цель работы" in s: sec = "goal"
        elif "2. Задание" in s: sec = "task"
        elif "3. Вывод" in s or "4. Вывод" in s: sec = "conclusion"
        elif "4. Результат" in s or "3. Результат" in s: sec = None
        else:
            if sec and s:
                if sec == "task":
                    data[sec] += s + " "
                else:
                    data[sec] += s + " "
    return data

def process_lab(lab_num):
    base_dir = r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии"
    source_path = os.path.join(base_dir, r"ЛР2\web-лр2_v2.docx")
    dest_path = os.path.join(base_dir, rf"ЛР{lab_num}\web-лр{lab_num}.docx")
    txt_path = os.path.join(base_dir, rf"ЛР{lab_num}\отчет.txt")
    photo_path = os.path.join(base_dir, rf"photo\lab{lab_num}_photo_1.png")
    
    if not os.path.exists(txt_path): return

    info = parse_txt(txt_path)
    doc = Document(source_path)

    # 1. Update Title Page (preserving formatting as much as possible)
    for p in doc.paragraphs:
        if "ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №2" in p.text:
            p.text = p.text.replace("ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №2", f"ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №{lab_num}")
        if "КАСКАДНЫЕ ТАБЛИЦЫ СТИЛЕЙ" in p.text:
            p.text = p.text.replace("КАСКАДНЫЕ ТАБЛИЦЫ СТИЛЕЙ (CSS)", info["title2"])
            p.text = p.text.replace("КАСКАДНЫЕ ТАБЛИЦЫ СТИЛЕЙ", info["title2"])
    
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if "ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №2" in p.text:
                        p.text = p.text.replace("ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №2", f"ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №{lab_num}")
                    if "КАСКАДНЫЕ ТАБЛИЦЫ СТИЛЕЙ" in p.text:
                        p.text = p.text.replace("КАСКАДНЫЕ ТАБЛИЦЫ СТИЛЕЙ (CSS)", info["title2"])
                        p.text = p.text.replace("КАСКАДНЫЕ ТАБЛИЦЫ СТИЛЕЙ", info["title2"])

    # 2. Find Sections
    paras = list(doc.paragraphs)
    goal_i, task_i, res_i, conc_i, app_i = -1, -1, -1, -1, -1
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t.startswith("1 Цель"): goal_i = i
        elif t.startswith("2 Задание"): task_i = i
        elif t.startswith("3 Результат"): res_i = i
        elif t.startswith("4 Вывод"): conc_i = i
        elif t.startswith("ПРИЛОЖЕНИЕ"): app_i = i

    # Helper to set content without stripping Bold from the header itself
    def set_content(start_idx, end_idx, content_text):
        if start_idx != -1 and end_idx != -1:
            # Clear everything between headers
            for j in range(start_idx + 1, end_idx):
                doc.paragraphs[j].text = ""
            # Set text in the first paragraph after header
            doc.paragraphs[start_idx + 1].text = content_text.strip()

    # Fill sections (keeping the template order: Goal -> Task -> Result -> Conclusion)
    if goal_i != -1 and task_i != -1: set_content(goal_i, task_i, info["goal"])
    if task_i != -1 and res_i != -1: set_content(task_i, res_i, info["task"])
    if conc_i != -1:
        end_val = app_i if app_i != -1 else len(doc.paragraphs)
        set_content(conc_i, end_val, info["conclusion"])

    # Result section (usually just the image)
    if res_i != -1 and conc_i != -1:
        for j in range(res_i + 1, conc_i):
            doc.paragraphs[j].text = "" # Clear old stuff
        
        # Insert Image and Caption
        img_p = doc.paragraphs[res_i + 1]
        run = img_p.add_run()
        if os.path.exists(photo_path):
            run.add_picture(photo_path, width=Inches(6.0))
        doc.paragraphs[res_i + 2].text = f"Рисунок 1 – Результат выполнения работы (ЛР{lab_num})"

    # 3. Appendix
    if app_i != -1:
        # Clear everything from App header to end
        for p in paras[app_i+1:]:
            p_xml = p._element
            parent = p_xml.getparent()
            if parent is not None: parent.remove(p_xml)
        
        # Re-add code
        lab_path = os.path.join(base_dir, f"ЛР{lab_num}")
        files = sorted(glob.glob(os.path.join(lab_path, "*.css"))) + sorted(glob.glob(os.path.join(lab_path, "*.html")))
        for fpath in files:
            doc.add_paragraph().add_run(f"\nКод файла {os.path.basename(fpath)}:")
            with open(fpath, 'r', encoding='utf-8') as f:
                p = doc.add_paragraph(f.read())
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(9)

    doc.save(dest_path)
    print(f"Done Lab {lab_num}")

for i in [3, 4, 5, 6]:
    process_lab(i)
