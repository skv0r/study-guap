from docx import Document
from docx.shared import Pt, Inches
import os

source_path = r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР1\web-лр1.docx"
dest_path = r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР2\web-лр2_v2.docx"
photo_path = r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\photo\lab2_photo_1.png"

def replace_in_paragraphs(paragraphs):
    for p in paragraphs:
        if "ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №1" in p.text:
            p.text = p.text.replace("ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №1", "ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №2")
        if "СТРУКТУРИРОВАНИЕ ТЕКСТА И СОЗДАНИЕ ГИПЕРССЫЛОК В WEB-ДОКУМЕНТЕ" in p.text:
            p.text = p.text.replace("СТРУКТУРИРОВАНИЕ ТЕКСТА И СОЗДАНИЕ ГИПЕРССЫЛОК В WEB-ДОКУМЕНТЕ", "КАСКАДНЫЕ ТАБЛИЦЫ СТИЛЕЙ (CSS)")

doc = Document(source_path)

# 1. Title/Header Replacement
replace_in_paragraphs(doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            replace_in_paragraphs(cell.paragraphs)

# 2. Section Replacement
paras = list(doc.paragraphs)
goal_start, task_start, result_start, conclusion_start, app_start = -1, -1, -1, -1, -1
for i, p in enumerate(paras):
    t = p.text.strip()
    if t.startswith("1 Цель"): goal_start = i
    elif t.startswith("2 Задание"): task_start = i
    elif t.startswith("3 Результат"): result_start = i
    elif t.startswith("4 Вывод"): conclusion_start = i
    elif t.startswith("ПРИЛОЖЕНИЕ А"): app_start = i

# Goal
if goal_start != -1 and task_start != -1:
    for i in range(goal_start + 1, task_start):
        p = paras[i]
        p.text = ""
    paras[goal_start + 1].text = "Изучение технологий каскадных таблиц стилей (CSS) для визуального форматирования элементов web-страниц. Применение внешних, внутренних и встроенных стилей, работа со шрифтами, отступами и псевдоклассами."

# Task
if task_start != -1 and result_start != -1:
    task_text = "На базе файлов первой лабораторной работы создать внешнюю таблицу стилей. Задать базовые настройки шрифтов, отступов и цветов. Подключить созданную внешнюю таблицу к разработанным ранее страницам. Применить внутренние стили в секции стиля, а также продемонстрировать использование встроенных стилей через одноименный атрибут. Настроить псевдоклассы для гиперссылок, задать рамки и внутренние отступы для блочных элементов. Настроить межстрочный интервал и внешние поля."
    for i in range(task_start + 1, result_start):
        paras[i].text = "" # clear old
    if task_start + 1 < result_start:
        paras[task_start + 1].text = task_text

# Conclusion
if conclusion_start != -1 and app_start != -1:
    for i in range(conclusion_start + 1, app_start):
        p = paras[i]
        p.text = ""
    paras[conclusion_start + 1].text = "В ходе лабораторной работы были получены навыки использования CSS для визуального оформления web-документов. На практике изучены тег link, блок style и атрибут style. Освоены свойства шрифта, отступы, рамки, а также псевдоклассы (:hover). Цель достигнута."

# 3. Images and Results Replace
fig1_idx, fig2_idx = -1, -1
for i, p in enumerate(paras):
    if "Рисунок 1" in p.text: fig1_idx = i
    if "Рисунок 2" in p.text: fig2_idx = i
    if "На рисунке 1-2 показан результат работы" in p.text:
        p.text = p.text.replace("На рисунке 1-2 показан результат работы", "На рисунке 1 показан результат работы")

# Replace Photo 1
if fig1_idx > 0:
    paras[fig1_idx].text = "Рисунок 1 – Страница с примененными стилями"
    p_img = paras[fig1_idx - 1]
    p_img.clear()
    run = p_img.add_run()
    if os.path.exists(photo_path):
        run.add_picture(photo_path, width=Inches(6.0))

# Remove Photo 2
if fig2_idx > 0:
    # Delete paragraph with caption and paragraph before it (image)
    for p in [paras[fig2_idx-1], paras[fig2_idx]]:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)

# 4. Appendix Replace
code_header_idx = -1
for i in range(app_start, len(paras)):
    if "Код страницы index.html" in paras[i].text or "В данном разделе представлен исходный код" in paras[i].text:
        code_header_idx = i
        break

if code_header_idx != -1:
    # Delete from 'В данном разделе...' and below
    for p in paras[code_header_idx:]:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
elif app_start != -1: # fallback
    for p in paras[app_start+1:]:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)

def append_code(title, path):
    if not os.path.exists(path): return
    doc.add_paragraph().add_run(f"\n{title}:")
    with open(path, 'r', encoding='utf-8') as f:
        p = doc.add_paragraph(f.read())
        p.runs[0].font.name = 'Courier New'
        p.runs[0].font.size = Pt(9)

append_code("Код файла style.css", r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР2\style.css")
append_code("Код файла index.html", r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР2\index.html")
append_code("Код файла page2.html", r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР2\page2.html")

# Remove extra empty lines around images for clean layout, optionally.
doc.save(dest_path)
print("SUCCESS!")
