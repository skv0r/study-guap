from docx import Document
from docx.shared import Pt, Inches
import os

def main():
    source_path = r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР1\web-лр1.docx"
    dest_path = r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР2\web-лр2.docx"
    photo_path = r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\photo\lab2_photo_1.png"
    
    doc = Document(source_path)
    
    replacements = {
        "ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №1": "ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №2",
        "ОФОРМЛЕНИЕ ТЕКСТА И СОЗДАНИЕ ГИПЕРССЫЛОК В WEB-ДОКУМЕНТЕ": "КАСКАДНЫЕ ТАБЛИЦЫ СТИЛЕЙ (CSS)",
        
        "Изучение HTML-тегов для форматирования текста и создания гиперссылок в web-документах. Ознакомление с основными элементами структуры HTML-страницы.": 
        "Изучение технологий каскадных таблиц стилей (CSS) для визуального форматирования элементов web-страниц.",
        
        "1.1. Создать первую web-страницу, на которой разместить": "1.1. На базе файлов ЛР1 создать внешнюю таблицу стилей style.css.",
        "короткий текст о себе;": "Задать базовые настройки шрифтов, отступов и цветов.",
        "горизонтальную линию;": "1.2. Подключить внешнюю таблицу к страницам.",
        "маркированный список с перечислением любимых стихов;": "1.3. Применить внутренние стили в секции стиль.",
        "гиперссылки на вторую страницу с текстами стихов, на скачивание файла и на адрес электронной почты.": "1.4. Продемонстрировать встроенные стили через атрибут style.",
        "1.2. Создать вторую web-страницу, на которой разместить": "1.5. Настроить псевдоклассы (:hover) для ссылок.",
        "заголовок и тексты трех любимых стихов разных авторов;": "1.6. Задать рамки (border) и внутренние отступы (padding) для блоков.",
        "графическое изображение-ссылку для возврата на первую страницу.": "1.7. Настроить межстрочный интервал (line-height) и поля (margin).",
        
        "В ходе лабораторной работы были освоены базовые принципы создания структуры HTML-документов. Изучены и успешно применены на практике теги для форматирования текста (заголовки, абзацы, списки), вставки независимых элементов оформления (горизонтальная линия), а также теги для вставки изображений. Особое внимание было уделено созданию различных видов гиперссылок: переходов между страницами одного сайта, ссылок на скачивание файлов и ссылок-вызовов почтового клиента. Установлено, что правильное использование атрибутов тегов (например, href, src, width) является критически важным для корректного отображения и функционирования элементов на web-странице. Поставленная цель работы достигнута в полном объеме.":
        "В ходе лабораторной работы были получены навыки использования каскадных таблиц стилей (CSS). Изучены способы внедрения стилей: тег link, блок style и атрибут style. Освоены свойства шрифта, внешние и внутренние отступы, рамки, а также псевдоклассы (:hover). Цель достигнута.",
        
        "Рисунок 1 – Страница сайта с разработанной структурой и гиперссылками": "Рисунок 1 – Внешний вид страницы с примененными CSS-стилями",
        "Код страницы index.html:": "Код страницы index.html:"
    }

    code_start_index = -1
    image_para_index = -1
    for i, p in enumerate(doc.paragraphs):
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, new)
        
        if "Рисунок 1" in p.text:
            image_para_index = i - 1

        if "Код страницы index.html" in p.text:
            code_start_index = i

    if image_para_index != -1 and image_para_index >= 0:
        p_img = doc.paragraphs[image_para_index]
        p_img.clear()
        run = p_img.add_run()
        if os.path.exists(photo_path):
            run.add_picture(photo_path, width=Inches(6.0))
            
    if code_start_index != -1:
        paras = list(doc.paragraphs)
        for p in paras[code_start_index+1:]:
            p._element.getparent().remove(p._element)
            
    def append_code(title, path):
        if not os.path.exists(path): return
        doc.add_paragraph().add_run(f"\n{title}:")
        with open(path, 'r', encoding='utf-8') as f:
            p = doc.add_paragraph(f.read())
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(9)
            
    append_code("Код файла style.css", r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР2\style.css")
    append_code("Код файла index.html", r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР2\index.html")
    append_code("Код файла page2.html", r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР2\page2.html")
    
    doc.save(dest_path)
    print(f"Successfully generated {dest_path}")

if __name__ == "__main__":
    main()
