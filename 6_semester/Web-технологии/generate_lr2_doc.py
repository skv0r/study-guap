import glob
import re
from docx import Document
from docx.shared import Pt
import os

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        # Simple replace for identical formatting
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

def main():
    source_path = r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР1\web-лр1.docx"
    dest_path = r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР2\web-лр2.docx"
    
    doc = Document(source_path)
    
    # Text replacements mapping
    replacements = {
        "ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №1": "ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №2",
        "ОФОРМЛЕНИЕ ТЕКСТА И СОЗДАНИЕ ГИПЕРССЫЛОК В WEB-ДОКУМЕНТЕ": "КАСКАДНЫЕ ТАБЛИЦЫ СТИЛЕЙ (CSS)",
        
        "Изучение HTML-тегов для форматирования текста и создания гиперссылок в web-документах. Ознакомление с основными элементами структуры HTML-страницы.": 
        "Изучение технологий каскадных таблиц стилей (CSS) для визуального форматирования элементов web-страниц. Ознакомление со способами подключения стилей (внешние, внутренние, встроенные), работы с цветом, шрифтами, рамками и псевдоклассами.",
        
        "1.1. Создать первую web-страницу, на которой разместить": "1.1. На базе файлов ЛР1 создать внешнюю таблицу стилей style.css",
        "короткий текст о себе;": "Задать базовые настройки шрифтов, отступов и цветов.",
        "горизонтальную линию;": "1.2. Подключить внешнюю таблицу к страницам.",
        "маркированный список с перечислением любимых стихов;": "1.3. Продемонстрировать внутренние стили в секции <style>.",
        "гиперссылки на вторую страницу с текстами стихов, на скачивание файла и на адрес электронной почты.": "1.4. Продемонстрировать встроенные стили через атрибут style.",
        "1.2. Создать вторую web-страницу, на которой разместить": "1.5. Настроить псевдоклассы (:hover) для ссылок.",
        "заголовок и тексты трех любимых стихов разных авторов;": "1.6. Задать рамки (border) и внутренние отступы (padding) для блоков.",
        "графическое изображение-ссылку для возврата на первую страницу.": "1.7. Настроить межстрочный интервал (line-height) и поля (margin).",
        
        "В ходе лабораторной работы были освоены базовые принципы создания структуры HTML-документов. Изучены и успешно применены на практике теги для форматирования текста (заголовки, абзацы, списки), вставки независимых элементов оформления (горизонтальная линия), а также теги для вставки изображений. Особое внимание было уделено созданию различных видов гиперссылок: переходов между страницами одного сайта, ссылок на скачивание файлов и ссылок-вызовов почтового клиента. Установлено, что правильное использование атрибутов тегов (например, href, src, width) является критически важным для корректного отображения и функционирования элементов на web-странице. Поставленная цель работы достигнута в полном объеме.":
        "В ходе лабораторной работы были получены навыки использования каскадных таблиц стилей (CSS) для визуального оформления web-документов. На практике изучены все три способа внедрения стилей: использование тега <link> для подключения внешнего файла style.css, блока <style> для настройки общих правил страницы и атрибута style для локального декорирования отдельных элементов. Освоены базовые свойства для управления параметрами шрифта. Продемонстрирована эффективная работа с блочной моделью: настройка внешних (margin) и внутренних (padding) отступов, а также применение различных типов рамок (border). Дополнительно изучены псевдоклассы (:hover) для создания эффектов интерактивности при наведении курсора на гиперссылки и изображения. Поставленная цель работы по освоению таблиц стилей достигнута в полном объеме.",
        
        "Рисунок 1 – Страница сайта с разработанной структурой и гиперссылками": "Рисунок 1 – Внешний вид страницы с примененными CSS-стилями",
        "Код страницы index.html:": "Код страницы index.html и style.css:"
    }

    # Find the paragraph saying "Код страницы index.html:"
    # Everything after this is code and needs to be replaced.
    code_start_index = -1
    for i, p in enumerate(doc.paragraphs):
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, new)
        
        if "Код страницы index.html и style.css:" in p.text:
            code_start_index = i
            break
            
    # Delete everything after the code header
    if code_start_index != -1:
        # Keep paragraphs up to code_start_index, delete the rest
        for p in doc.paragraphs[code_start_index+1:]:
            p_element = p._element
            p_element.getparent().remove(p_element)
            
    # Now read the ACTUAL Lab 2 code and dump it in!
    def append_code(title, path):
        if not os.path.exists(path): return
        doc.add_paragraph().add_run(f"\n{title}:")
        with open(path, 'r', encoding='utf-8') as f:
            code_text = f.read()
            p = doc.add_paragraph(code_text)
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(9)
            
    append_code("Код файла style.css", r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР2\style.css")
    append_code("Код файла index.html", r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР2\index.html")
    append_code("Код файла page2.html", r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии\ЛР2\page2.html")
    
    # Save the modified document
    doc.save(dest_path)
    print(f"Successfully generated {dest_path}")

if __name__ == "__main__":
    main()
