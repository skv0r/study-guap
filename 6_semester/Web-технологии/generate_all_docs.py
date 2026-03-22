from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK
import os
import glob

def parse_txt(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    d = {"title1": "", "title2": "", "goal": "", "task": "", "conclusion": ""}
    sec = None
    for l in lines:
        s = l.strip()
        if "ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ" in s: d["title1"] = s
        elif d["title1"] and not d["title2"] and s and not any(x in s for x in ["ВЫПОЛНИЛ", "ПРЕПОДАВАТЕЛЬ", "Санкт-Петербург", "42", "Дисциплина"]): 
            if not s.startswith("Р"): d["title2"] = s
                 
        if "1. Цель работы" in s: sec = "goal"
        elif "2. Задание" in s: sec = "task"
        elif "3. Вывод" in s: sec = "conclusion"
        elif "4. Результат" in s: sec = None
        else:
            if sec and s: d[sec] += (" " if d[sec] else "") + s
    return d

def generate_for_lab(lab_num):
    base_dir = r"c:\Users\skvor\OneDrive\Desktop\GitHub\study-guap\6_semester\Web-технологии"
    source_path = os.path.join(base_dir, r"ЛР2\web-лр2_v2.docx")
    dest_path = os.path.join(base_dir, rf"ЛР{lab_num}\web-лр{lab_num}.docx")
    txt_path = os.path.join(base_dir, rf"ЛР{lab_num}\отчет.txt")
    photo_path = os.path.join(base_dir, rf"photo\lab{lab_num}_photo_1.png")
    
    if not os.path.exists(txt_path):
        print(f"Skipping Lab {lab_num}, no отчет.txt")
        return
        
    data = parse_txt(txt_path)
    if not data["title2"]:
        titles = {3:"РАБОТА С ТАБЛИЦАМИ И УПРАВЛЕНИЕ ФОНОМ В WEB-ДОКУМЕНТЕ", 
                  4:"РАБОТА С ВЕБ-ФОРМАМИ", 
                  5:"ВЕРСТКА СТРАНИЦ С ИСПОЛЬЗОВАНИЕМ ТЕХНОЛОГИИ FLEXBOX", 
                  6:"ВЕРСТКА СТРАНИЦ С ИСПОЛЬЗОВАНИЕМ ТЕХНОЛОГИИ CSS GRID"}
        data["title2"] = titles.get(lab_num, "")

    doc = Document(source_path)
    
    def repl_title(paras):
        for p in paras:
            if "ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №2" in p.text:
                for run in p.runs:
                    if "ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №2" in run.text:
                        run.text = run.text.replace("ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №2", f"ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ №{lab_num}")
            if "КАСКАДНЫЕ ТАБЛИЦЫ СТИЛЕЙ" in p.text:
                for run in p.runs:
                    if "КАСКАДНЫЕ ТАБЛИЦЫ СТИЛЕЙ" in run.text:
                        # Find the whole sequence that looks like the title
                        run.text = run.text.replace("КАСКАДНЫЕ ТАБЛИЦЫ СТИЛЕЙ", data["title2"])
                        if "(CSS)" in run.text: run.text = run.text.replace("(CSS)", "")

    repl_title(doc.paragraphs)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                repl_title(c.paragraphs)
                
    paras = list(doc.paragraphs)
    goal_i, task_i, res_i, conc_i, app_i = -1,-1,-1,-1,-1
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t.startswith("1 Цель"): goal_i = i
        elif t.startswith("2 Задание"): task_i = i
        elif t.startswith("3 Результат"): res_i = i
        elif t.startswith("4 Вывод"): conc_i = i
        elif t.startswith("ПРИЛОЖЕНИЕ"): app_i = i

    def set_bold_text(para_idx, text):
        if para_idx != -1:
            p = doc.paragraphs[para_idx]
            p.text = ""
            run = p.add_run(text)
            run.bold = True

    # Lab 3-6 require: 1. Goal, 2. Task, 3. Conclusion, 4. Result
    # Template (Lab 2) has: 1. Goal, 2. Task, 3. Result, 4. Conclusion
    
    # Let's fix numbering and headers first
    indices = sorted([goal_i, task_i, res_i, conc_i])
    # idx 0: Goal
    # idx 1: Task
    # idx 2: Conclusion
    # idx 3: Result
    
    # We re-assign the headers to the sorted positions to ensure 3 is always Conclusion and 4 is always Result
    set_bold_text(indices[0], "1 Цель работы")
    set_bold_text(indices[1], "2 Задание")
    set_bold_text(indices[2], "3 Вывод")
    set_bold_text(indices[3], "4 Результат работы")
    
    # Update contents between them
    def fill_between(start_idx, end_idx, text):
        if start_idx != -1 and end_idx != -1:
            for i in range(start_idx + 1, end_idx):
                doc.paragraphs[i].text = ""
            doc.paragraphs[start_idx + 1].text = text

    fill_between(indices[0], indices[1], data["goal"])
    fill_between(indices[1], indices[2], data["task"])
    fill_between(indices[2], indices[3], data["conclusion"])

    # Result section (indices[3] to Appendix)
    result_end = app_i if app_i != -1 else len(doc.paragraphs)
    for i in range(indices[3] + 1, result_end):
        doc.paragraphs[i].text = ""

    # Page Breaks backwards
    all_headers = sorted([indices[1], indices[2], indices[3], app_i], reverse=True)
    for idx in all_headers:
        if idx != -1:
            doc.paragraphs[idx].insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)

    # Image replacement
    new_paras = list(doc.paragraphs)
    for i, p in enumerate(new_paras):
        if "4 Результат" in p.text:
            # The image should be in the paragraph AFTER "4 Результат" (which was indices[3]+1 originally)
            # Find the caption "Рисунок 1"
            for j in range(i + 1, len(new_paras)):
                if "Рисунок 1" in new_paras[j].text:
                    new_paras[j].text = f"Рисунок 1 – Результат страницы (ЛР{lab_num})"
                    # Bold the caption
                    new_paras[j].runs[0].bold = True
                    p_img = new_paras[j-1]
                    p_img.clear()
                    run = p_img.add_run()
                    if os.path.exists(photo_path):
                        run.add_picture(photo_path, width=Inches(6.0))
                    break
            break
            
    # Appendix cleanup
    new_app_i = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("ПРИЛОЖЕНИЕ"):
            new_app_i = i
            break
            
    if new_app_i != -1:
        for p in list(doc.paragraphs)[new_app_i+1:]:
            p_element = p._element
            parent = p_element.getparent()
            if parent is not None: parent.remove(p_element)

    def append_code(title, path):
        if not os.path.exists(path): return
        doc.add_paragraph().add_run(f"\n{title}:")
        with open(path, 'r', encoding='utf-8') as f:
            p = doc.add_paragraph(f.read())
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(9)
            
    lab_dir = os.path.join(base_dir, f"ЛР{lab_num}")
    css_files = sorted(glob.glob(os.path.join(lab_dir, "*.css")))
    html_files = sorted(glob.glob(os.path.join(lab_dir, "*.html")))
    for c in css_files: append_code(f"Код файла {os.path.basename(c)}", c)
    for h in html_files: append_code(f"Код файла {os.path.basename(h)}", h)
    
    doc.save(dest_path)
    print(f"Success Lab {lab_num}")

if __name__ == "__main__":
    for i in [3, 4, 5, 6]:
        generate_for_lab(i)
