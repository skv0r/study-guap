#!/usr/bin/env python3
"""Отчёт ЛР1 по ГОСТ 7.32 / ГОСТ 2.105 и бланку ГУАП (институты 1–4)."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

BASE = Path(__file__).resolve().parent
SHOTS = BASE / "screenshots"
BLANK = BASE / "guap_blanks" / "lab.docx"
OUT = BASE / "ЛР1_Отчет_Буренков_PairWise.docx"
PAIRWISE = Path("/Users/gregoryburenkov/dev/GitHub/pairwise")


def set_run_font(run, size=14, bold=False, name="Times New Roman", italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def format_paragraph(
    p,
    *,
    first_indent=True,
    align="justify",
    space_after=0,
    space_before=0,
    left_indent=0,
    line_spacing=1.5,
):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.first_line_indent = Cm(1.25 if first_indent else 0)
    pf.left_indent = Cm(left_indent)
    p.alignment = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]


def add_body(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p)
    set_run_font(p.add_run(text))
    return p


def add_section_h(doc, text, *, new_page=True):
    """Заголовок раздела: с абзацного отступа, с прописной, полужирный (ГОСТ)."""
    if new_page:
        doc.add_page_break()
    p = doc.add_paragraph()
    format_paragraph(p, first_indent=True, align="left", space_before=0, space_after=12)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), bold=True)


def add_sub_h(doc, text):
    """Заголовок подраздела: с абзацного отступа, полужирный."""
    p = doc.add_paragraph()
    format_paragraph(p, first_indent=True, align="left", space_before=12, space_after=6)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), bold=True)


def add_struct(doc, text, *, new_page=True):
    """Структурный элемент: по центру, ПРОПИСНЫМИ, без точки."""
    if new_page:
        doc.add_page_break()
    p = doc.add_paragraph()
    format_paragraph(p, first_indent=False, align="center", space_after=18)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), bold=True)


def add_toc_line(doc, title, page, *, sub=False):
    p = doc.add_paragraph()
    format_paragraph(p, first_indent=False, align="left", left_indent=0.5 if sub else 0)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    set_run_font(p.add_run(title))
    p.add_run("\t")
    set_run_font(p.add_run(str(page)))


def add_caption_table(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p, first_indent=False, align="left", space_before=10, space_after=4)
    set_run_font(p.add_run(text))


def add_figure(doc, path, caption, w=15.5):
    p = doc.add_paragraph()
    format_paragraph(p, first_indent=False, align="center", space_before=10)
    p.add_run().add_picture(str(path), width=Cm(w))
    c = doc.add_paragraph()
    format_paragraph(c, first_indent=False, align="center", space_before=6, space_after=10)
    set_run_font(c.add_run(caption))


def _set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for cell in row.cells:
            _set_cell_border(cell)
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        format_paragraph(p, first_indent=False, align="center")
        set_run_font(p.add_run(h), bold=True, size=12)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            format_paragraph(p, first_indent=False, align="left")
            set_run_font(p.add_run(v), size=12)
    spacer = doc.add_paragraph()
    format_paragraph(spacer, first_indent=False, line_spacing=1.0, space_after=6)


def add_listing(doc, title: str, code: str):
    p = doc.add_paragraph()
    format_paragraph(p, first_indent=False, align="left", space_before=10, space_after=4)
    set_run_font(p.add_run(title), bold=True)
    for line in code.strip("\n").splitlines():
        lp = doc.add_paragraph()
        format_paragraph(lp, first_indent=False, align="left", line_spacing=1.0, space_after=0)
        set_run_font(lp.add_run(line if line else " "), size=10, name="Courier New")


def add_page_field(paragraph):
    run = paragraph.add_run()
    for kind, val in [("begin", None), ("instr", " PAGE "), ("end", None)]:
        if kind == "instr":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = val
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        run._r.append(el)
    set_run_font(run)


def put_cell(cell, text, size=14):
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r.text = ""
    run = p.add_run(text)
    set_run_font(run, size=size)


def fill_title(doc: Document):
    for p in doc.paragraphs:
        if p.text.startswith("КАФЕДРА"):
            if len(p.runs) >= 2 and p.runs[1].text.strip("_") == "":
                p.runs[1].text = "42"
            else:
                p.clear()
                r = p.add_run("КАФЕДРА № 42")
                set_run_font(r, size=14)
        elif "20__" in p.text:
            for r in p.runs:
                if "20__" in r.text:
                    r.text = r.text.replace("20__", "2026")

    t_prep, t_work, t_stud = doc.tables
    put_cell(t_prep.rows[0].cells[0], "ассистент")
    put_cell(t_prep.rows[0].cells[4], "В.В. Жукалин")

    put_cell(t_work.rows[0].cells[0], "ОТЧЕТ О ЛАБОРАТОРНОЙ РАБОТЕ № 1")
    put_cell(
        t_work.rows[1].cells[0],
        "Создание проекта Next.js с использованием React, TypeScript и App Router",
    )
    put_cell(t_work.rows[2].cells[0], "по курсу: Web-программирование")

    put_cell(t_stud.rows[0].cells[1], "4321")
    put_cell(t_stud.rows[0].cells[5], "Г.В. Буренков")


def setup_body_section(doc: Document):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    title_sec, body_sec = doc.sections[0], doc.sections[1]

    title_sec.footer.is_linked_to_previous = False
    for p in title_sec.footer.paragraphs:
        p.clear()

    body_sec.page_width, body_sec.page_height = Mm(210), Mm(297)
    body_sec.left_margin, body_sec.right_margin = Mm(30), Mm(15)
    body_sec.top_margin, body_sec.bottom_margin = Mm(20), Mm(20)
    body_sec.footer.is_linked_to_previous = False
    body_sec.header.is_linked_to_previous = False

    sect_pr = body_sec._sectPr
    for old in sect_pr.findall(qn("w:pgNumType")):
        sect_pr.remove(old)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), "2")
    sect_pr.append(pg)

    fp = body_sec.footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)


def read_src(rel: str) -> str:
    return (PAIRWISE / rel).read_text(encoding="utf-8")


def build_body(doc: Document):
    add_struct(doc, "СОДЕРЖАНИЕ", new_page=False)
    add_toc_line(doc, "1 Цель работы", 3)
    add_toc_line(doc, "2 Концепция приложения", 4)
    add_toc_line(doc, "2.1 Проблема и целевая аудитория", 4, sub=True)
    add_toc_line(doc, "2.2 Главная предметная сущность", 4, sub=True)
    add_toc_line(doc, "3 Подготовка среды разработки", 5)
    add_toc_line(doc, "3.1 Версии Node.js и npm", 5, sub=True)
    add_toc_line(doc, "3.2 Создание проекта", 5, sub=True)
    add_toc_line(doc, "4 Структура и реализация", 6)
    add_toc_line(doc, "4.1 Структура каталогов", 6, sub=True)
    add_toc_line(doc, "4.2 Компонент FeatureCard", 6, sub=True)
    add_toc_line(doc, "4.3 Страницы приложения", 7, sub=True)
    add_toc_line(doc, "5 Результаты проверки", 8)
    add_toc_line(doc, "5.1 Внешний вид страниц", 8, sub=True)
    add_toc_line(doc, "5.2 Сборка проекта", 9, sub=True)
    add_toc_line(doc, "6 Ответы на контрольные вопросы", 10)
    add_toc_line(doc, "ЗАКЛЮЧЕНИЕ", 12)
    add_toc_line(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 13)
    add_toc_line(doc, "ПРИЛОЖЕНИЕ А Листинги исходного кода", 14)

    add_section_h(doc, "1 Цель работы")
    add_body(
        doc,
        "Целью работы является подготовка среды разработки и создание начальной "
        "версии веб-приложения на основе Next.js, React и TypeScript с "
        "использованием App Router [1], [2]. В течение семи лабораторных работ "
        "разрабатывается одно пользовательское приложение; в данной работе "
        "формируется его каркас: главная страница, информационная страница /about, "
        "типизированный переиспользуемый компонент и навигация между маршрутами.",
    )

    add_section_h(doc, "2 Концепция приложения")
    add_body(
        doc,
        "Рабочее название приложения — PairWise. Сервис предназначен для поиска "
        "коворкинг-зон, в которых удобно проводить очные встречи с "
        "единомышленниками: разобрать рабочие задачи, обсудить тренды, найти "
        "людей со схожим стеком и проектами. Тематика не копирует учебный пример "
        "TheraCity: предметом является не фотография и не прогулочный маршрут, "
        "а площадка для совместной работы.",
    )
    add_body(
        doc,
        "К седьмой лабораторной работе PairWise должен стать законченным "
        "приложением с несколькими связанными разделами: каталог зон, карточка "
        "места, карта, избранное, личный кабинет, страницы входа и регистрации. "
        "В первой работе реализован только начальный каркас: маршруты / и /about, "
        "корневой layout и компонент FeatureCard. Исходный код страниц и "
        "компонента приведён в приложении А.",
    )

    add_sub_h(doc, "2.1 Проблема и целевая аудитория")
    add_body(
        doc,
        "Проблема состоит в том, что разработчикам и другим специалистам из "
        "IT-сферы сложно быстро найти место для очной встречи, где будут люди "
        "со схожими интересами, стеком и проектами. Сведения о коворкингах "
        "разрознены, а фильтров по сообществу площадки обычно нет. PairWise "
        "должен собрать площадки в одном каталоге и помочь выбрать зону под "
        "конкретный формат встречи.",
    )
    add_body(
        doc,
        "Целевая аудитория — разработчики, фаундеры и люди из IT-сферы, которым "
        "нужны очные встречи по делу, а не только удалённые созвоны.",
    )

    add_sub_h(doc, "2.2 Главная предметная сущность")
    add_body(
        doc,
        "Главная предметная сущность — коворкинг-зона: конкретная площадка, "
        "куда можно прийти на очную встречу. У зоны есть описание, стек и "
        "интересы сообщества; позднее появятся адрес, фотографии и точка на "
        "карте. Каталог, карточка места, избранное и совместные визиты "
        "строятся вокруг этого объекта. На данном этапе сущность описана на "
        "страницах приложения и задаёт направление следующих лабораторных работ.",
    )

    add_section_h(doc, "3 Подготовка среды разработки")
    add_sub_h(doc, "3.1 Версии Node.js и npm")
    add_body(
        doc,
        "Работа выполнялась в операционной системе macOS со стандартным "
        "приложением «Терминал». Согласно документации Next.js минимальная "
        "поддерживаемая версия Node.js составляет 20.9 [2]. После установки "
        "проверены версии интерпретатора и менеджера пакетов: Node.js 20.18.0, "
        "npm 10.8.2. Полученная версия Node.js удовлетворяет требованию курса.",
    )

    add_sub_h(doc, "3.2 Создание проекта")
    add_body(
        doc,
        "Проект создан генератором create-next-app. Next.js, React, React DOM "
        "и TypeScript устанавливаются автоматически и не требуют отдельной "
        "установки [2]. Название учебного каталога student-project заменено "
        "на pairwise:",
    )
    add_listing(doc, "Команда создания проекта", "npx create-next-app@latest pairwise")
    add_body(
        doc,
        "Выбрана пользовательская настройка генератора. Обязательные параметры "
        "курса — TypeScript, App Router и ESLint — включены [3], [7]. Каталог "
        "src не использовался: каталог app расположен в корне проекта. "
        "Выбранные параметры приведены в таблице 1.",
    )
    add_caption_table(doc, "Таблица 1 — Параметры create-next-app")
    add_table(
        doc,
        ["Параметр", "Значение"],
        [
            ["TypeScript", "Yes"],
            ["Linter", "ESLint"],
            ["React Compiler", "No"],
            ["Tailwind CSS", "No"],
            ["src directory", "No"],
            ["App Router", "Yes"],
            ["Import alias", "@/*"],
        ],
    )
    add_body(
        doc,
        "После генерации в файле package.json зафиксированы зависимости "
        "next 16.3.4, react 19.2.8 и react-dom 19.2.8. Назначение стандартных "
        "сценариев npm приведено в таблице 2.",
    )
    add_caption_table(doc, "Таблица 2 — Сценарии package.json")
    add_table(
        doc,
        ["Команда", "Назначение"],
        [
            ["npm run dev", "запуск сервера разработки"],
            ["npm run build", "производственная сборка"],
            ["npm run start", "запуск ранее собранного приложения"],
            ["npm run lint", "статическая проверка ESLint"],
        ],
    )

    add_section_h(doc, "4 Структура и реализация")
    add_sub_h(doc, "4.1 Структура каталогов")
    add_body(
        doc,
        "Маршруты App Router размещаются в каталоге app: каталоги задают "
        "сегменты URL, специальные файлы — содержимое маршрутов [4], [5]. "
        "Файл app/page.tsx соответствует адресу /, файл app/about/page.tsx — "
        "адресу /about. Каталог ui не является специальным каталогом Next.js "
        "и используется в курсе для хранения переиспользуемых компонентов. "
        "Итоговая структура проекта приведена на рисунке 1.",
    )
    add_figure(doc, SHOTS / "structure.png", "Рисунок 1 — Структура каталогов проекта PairWise", w=14.0)
    add_body(
        doc,
        "Корневой layout.tsx задаёт общую HTML-оболочку, язык интерфейса ru "
        "и метаданные. Содержимое текущей страницы передаётся через свойство "
        "children [5]. Стили страниц и компонента вынесены в CSS-модули.",
    )

    add_sub_h(doc, "4.2 Компонент FeatureCard")
    add_body(
        doc,
        "Создан переиспользуемый React-компонент FeatureCard [8]. Компонент "
        "принимает название функции и её краткое описание через свойства "
        "title и description. Свойства типизированы и доступны только для "
        "чтения [10], [12]. Компонент возвращает JSX-разметку карточки [9] "
        "и используется на главной странице три раза. Код компонента приведён "
        "в листинге А.1 приложения А.",
    )

    add_sub_h(doc, "4.3 Страницы приложения")
    add_body(
        doc,
        "Главная страница содержит рабочее название PairWise, описание "
        "решаемой проблемы, целевую аудиторию, три предполагаемые функции "
        "в виде компонентов FeatureCard и ссылку на страницу /about. Для "
        "перехода используется компонент Link из next/link. Код главной "
        "страницы приведён в листинге А.2.",
    )
    add_body(
        doc,
        "Страница /about содержит цель разработки, описание будущего "
        "результата, главную предметную сущность, предполагаемые разделы "
        "приложения и ссылку для возврата на главную страницу. Код страницы "
        "приведён в листинге А.3, код корневого layout — в листинге А.4. "
        "Директива «use client» не применялась: страницы статические, без "
        "состояния и обращения к браузерным API [6].",
    )

    add_section_h(doc, "5 Результаты проверки")
    add_sub_h(doc, "5.1 Внешний вид страниц")
    add_body(
        doc,
        "Сервер разработки запускается командой npm run dev. Главная страница "
        "доступна по адресу http://localhost:3000/, информационная — по адресу "
        "http://localhost:3000/about. Внешний вид страниц приведён на "
        "рисунках 2 и 3. Навигация выполняется ссылками «О нас» и "
        "«Вернуться на главную страницу».",
    )
    add_figure(
        doc,
        SHOTS / "home.png",
        "Рисунок 2 — Главная страница PairWise",
        w=15.5,
    )
    add_figure(
        doc,
        SHOTS / "about.png",
        "Рисунок 3 — Страница «О проекте»",
        w=15.5,
    )

    add_sub_h(doc, "5.2 Сборка проекта")
    add_body(
        doc,
        "Команда npm run lint завершилась без сообщений об ошибках. "
        "Производственная сборка npm run build выполнена успешно: Next.js "
        "сообщил Compiled successfully, проверка TypeScript прошла без ошибок, "
        "маршруты / и /about сформированы как статические страницы. Фрагмент "
        "протокола сборки приведён ниже.",
    )
    add_listing(
        doc,
        "Фрагмент результата npm run build",
        """\
▲ Next.js 16.3.4 (Turbopack)
✓ Compiled successfully in 3.3s
  Running TypeScript ...
  Finished TypeScript in 1418ms ...

Route (app)
┌ ○ /
├ ○ /_not-found
└ ○ /about

○  (Static)  prerendered as static content""",
    )
    add_body(
        doc,
        "Проект считается технически корректным. Серверная часть, база данных "
        "и авторизация в этой работе не добавлялись и будут реализованы в "
        "следующих лабораторных работах.",
    )

    add_section_h(doc, "6 Ответы на контрольные вопросы")
    questions = [
        (
            "Что такое React-компонент и какие требования предъявляются к его имени?",
            "React-компонент — функция, которая возвращает разметку "
            "пользовательского интерфейса. Компоненты позволяют разделять "
            "интерфейс на самостоятельные части и повторно использовать их [8]. "
            "Имя компонента должно начинаться с заглавной буквы. В работе таким "
            "компонентом является FeatureCard.",
        ),
        (
            "Что такое JSX и чем он отличается от HTML?",
            "JSX — расширение синтаксиса JavaScript, позволяющее описывать "
            "разметку внутри кода [9]. Он похож на HTML, но строже: компонент "
            "должен возвращать один корневой элемент, все теги закрываются, "
            "а JavaScript-выражения вставляются в фигурных скобках.",
        ),
        (
            "Как данные передаются в дочерний компонент?",
            "Данные передаются через свойства (props), аналогично HTML-атрибутам, "
            "и доступны только для чтения [10]. В FeatureCard свойства title и "
            "description типизированы, чтобы TypeScript выявлял ошибки до запуска "
            "программы [11], [12].",
        ),
        (
            "Для чего в проекте используется TypeScript?",
            "TypeScript расширяет JavaScript системой типов и помогает находить "
            "неожиданное поведение до выполнения программы [11]. Файлы без JSX "
            "имеют расширение .ts, файлы с JSX — .tsx, параметры компилятора "
            "задаются в tsconfig.json [7].",
        ),
        (
            "Что добавляет Next.js к React?",
            "React отвечает за компонентную модель интерфейса. Next.js добавляет "
            "структуру приложения, файловую маршрутизацию, серверный рендеринг "
            "и стандартные команды разработки [1], [3]. Рекомендуемый способ "
            "создания проекта — create-next-app [2].",
        ),
        (
            "Как App Router связывает каталоги с адресами URL?",
            "Маршруты размещаются в каталоге app. Каталоги задают сегменты URL, "
            "специальные файлы — содержимое маршрута [3], [4]. Файл app/page.tsx "
            "соответствует адресу /, файл app/about/page.tsx — адресу /about. "
            "Компонент страницы экспортируется по умолчанию [5].",
        ),
        (
            "Чем отличаются page.tsx и layout.tsx?",
            "Файл page.tsx определяет страницу по конкретному адресу. Файл "
            "layout.tsx задаёт общую разметку для нескольких страниц. Корневой "
            "layout обязан содержать элементы html и body; содержимое страницы "
            "передаётся через children [5].",
        ),
        (
            "Когда нужен клиентский компонент и директива «use client»?",
            "По умолчанию страницы и layouts в App Router являются Server "
            "Components. Клиентский компонент нужен при обработчиках событий, "
            "состоянии, эффектах и обращении к браузерным API [6]. Директива "
            "«use client» ставится в начале файла. В первой работе она не "
            "требуется: создан статический каркас.",
        ),
        (
            "Какие команды проверяют работоспособность проекта?",
            "Команда npm run dev запускает сервер разработки, npm run lint "
            "выполняет статическую проверку, npm run build создаёт "
            "производственную сборку, npm run start запускает собранное "
            "приложение [2]. Проект корректен, если lint и build завершаются "
            "без ошибок.",
        ),
    ]
    for i, (q, a) in enumerate(questions, 1):
        add_body(doc, f"{i}. {q} {a}")

    add_struct(doc, "ЗАКЛЮЧЕНИЕ")
    add_body(
        doc,
        "В ходе лабораторной работы подготовлена среда разработки и создан "
        "проект PairWise на Next.js, React и TypeScript с App Router. "
        "Сформулирована концепция приложения: поиск коворкинг-зон для очных "
        "встреч с единомышленниками. Главная предметная сущность — коворкинг-зона.",
    )
    add_body(
        doc,
        "Реализованы главная страница и страница /about, типизированный "
        "компонент FeatureCard используется три раза, навигация выполняется "
        "через next/link. Проект проходит команды npm run lint и npm run build "
        "без ошибок. Серверная часть, база данных и авторизация не добавлялись. "
        "Цель работы достигнута; полученный каркас будет развиваться в "
        "следующих лабораторных работах.",
    )

    add_struct(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    sources = [
        "Next.js : Getting Started. – URL: https://nextjs.org/docs/getting-started (дата обращения: 02.09.2026).",
        "Next.js : Installation. – URL: https://nextjs.org/docs/app/getting-started/installation (дата обращения: 02.09.2026).",
        "Next.js : App Router. – URL: https://nextjs.org/docs/app (дата обращения: 02.09.2026).",
        "Next.js : Project Structure. – URL: https://nextjs.org/docs/app/getting-started/project-structure (дата обращения: 02.09.2026).",
        "Next.js : Layouts and Pages. – URL: https://nextjs.org/docs/app/building-your-application/routing/pages-and-layouts (дата обращения: 02.09.2026).",
        "Next.js : Server and Client Components. – URL: https://nextjs.org/docs/app/building-your-application/rendering/composition-patterns (дата обращения: 02.09.2026).",
        "Next.js : TypeScript. – URL: https://nextjs.org/docs/app/building-your-application/configuring/typescript (дата обращения: 02.09.2026).",
        "React : Your First Component. – URL: https://react.dev/learn/your-first-component (дата обращения: 02.09.2026).",
        "React : Writing Markup with JSX. – URL: https://react.dev/learn/writing-markup-with-jsx (дата обращения: 02.09.2026).",
        "React : Passing Props to a Component. – URL: https://react.dev/learn/passing-props-to-a-component (дата обращения: 02.09.2026).",
        "TypeScript for JavaScript Programmers. – URL: https://www.typescriptlang.org/docs/handbook/typescript-in-5-minutes.html (дата обращения: 02.09.2026).",
        "TypeScript : Everyday Types. – URL: https://www.typescriptlang.org/docs/handbook/2/everyday-types.html (дата обращения: 02.09.2026).",
        "ГОСТ 7.32-2017. Отчет о научно-исследовательской работе. Структура и правила оформления. – М. : Стандартинформ, 2017.",
        "ГОСТ 2.105-2019. Общие требования к текстовым документам. – М. : Стандартинформ, 2019.",
        "ГОСТ Р 7.0.100-2018. Библиографическая запись. Библиографическое описание. Общие требования и правила составления. – М. : Стандартинформ, 2018.",
    ]
    for i, src in enumerate(sources, 1):
        p = doc.add_paragraph()
        format_paragraph(p)
        set_run_font(p.add_run(f"{i}. {src}"))

    add_struct(doc, "ПРИЛОЖЕНИЕ А")
    p = doc.add_paragraph()
    format_paragraph(p, first_indent=False, align="center", space_after=12)
    set_run_font(p.add_run("Листинги исходного кода"), bold=True)
    add_body(
        doc,
        "В приложении приведены исходные тексты компонента FeatureCard, "
        "главной страницы, страницы /about и корневого layout.",
    )
    add_listing(doc, "Листинг А.1 — Файл app/ui/FeatureCard.tsx", read_src("app/ui/FeatureCard.tsx"))
    add_listing(doc, "Листинг А.2 — Файл app/page.tsx", read_src("app/page.tsx"))
    add_listing(doc, "Листинг А.3 — Файл app/about/page.tsx", read_src("app/about/page.tsx"))
    add_listing(doc, "Листинг А.4 — Файл app/layout.tsx", read_src("app/layout.tsx"))


def main():
    shutil.copy(BLANK, OUT)
    doc = Document(str(OUT))
    fill_title(doc)
    setup_body_section(doc)
    build_body(doc)
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
